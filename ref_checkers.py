"""
This module aims at checking the references in the document.
It looks whether the inline citations in the document were correctly
given by the person in the reference list.

NOTE: This is only for apa style, we'll have to also check for the MLA style
"""
import re
import docx

# open the document
doc = docx.Document("./docx_files/data_viz.docx")

# extract the text
doc_text = "\n".join([para.text for para in doc.paragraphs])


def get_inline_citations_dict() -> dict:
    """
    Get all inline citations and return them as a key value pair
    where the key is the citation and the value is the text preceding
    the citation itself.
    """
    # define the regular expression to find the inline citations
    # pattern for text inside parentheses
    inline_citation_pattern = r"\(([^,]*,[^,]*,[^,]*)\)"

    # find all inline citations and their preceding sentences
    inline_citations: dict = {}

    references_heading_found = False

    for match in re.finditer(inline_citation_pattern, doc_text):
        if references_heading_found:
            break

        citation_text = match.group(1)
        citation_index = match.start(0)

        # find the start of the preceding sentence
        sentence_start = doc_text.rfind(".", 0, citation_index) + 1
        preceding_sentence = doc_text[sentence_start:citation_index].strip()
        inline_citations[citation_text] = preceding_sentence

        if "references" in preceding_sentence.lower():
            references_heading_found = True

    return inline_citations


def get_first_name_in_citation(citation: str) -> str:
    """
    Get the first name in the citation.
    """
    return citation.split(",")[0].strip().split(" ")[0]


def get_inline_citations_list() -> list:
    """
    Get all inline citations from the get_inline_citations_dict
    and return them as a list.
    """

    # if the citations very first letter can be converted to an int
    # then it is a number and not a name
    # so we can ignore it

    inline_citations: dict = get_inline_citations_dict()
    inline_citations_list: list = []

    for citation in inline_citations.keys():
        if not citation[0].isdigit():
            inline_citations_list.append(citation)

    return list(set(inline_citations_list))


def get_references_list() -> list:
    """
    Get all references in the docx file and return them as a list.
    """
    references_heading: None = None
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading") and\
                "references" in paragraph.text.lower():
            references_heading = paragraph
            break

    found_ref: bool = False
    references: list = []

    if references_heading:
        for paragraph in doc.paragraphs:
            if paragraph.text == references_heading.text:
                found_ref: bool = True
            elif found_ref:
                if paragraph.style.name.startswith("Heading"):
                    break
                references.append(paragraph.text)

    return references


def get_citations_and_references_dict() -> dict:
    """
    Get all citations and references and return them as a key value pair
    """

    matched_citations: dict = {}

    for citation in get_inline_citations_list():
        citation_name = get_first_name_in_citation(citation)
        matching_refs: list = []
        for ref in get_references_list():
            if citation_name.lower() in ref.lower():
                matching_refs.append(ref)
        matched_citations[citation] = matching_refs

    return matched_citations


def main() -> None:
    """Main function"""
    cit_and_ref = get_citations_and_references_dict()

    for citation, ref in cit_and_ref.items():
        print(citation, "\n\n", ref, "\n\n")


if __name__ == "__main__":
    main()
